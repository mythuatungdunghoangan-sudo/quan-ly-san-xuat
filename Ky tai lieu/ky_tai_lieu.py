"""
Ứng dụng ký tài liệu
- Upload 1 hoặc nhiều file (PDF/ảnh/Excel/Word), tự động tìm vị trí ký theo từ khóa
- Hỗ trợ AI Vision (Claude API) tìm vị trí khi không quét được, xuất ZIP
"""

import io, zipfile
import numpy as np
import streamlit as st
from PIL import Image, ImageDraw
from pathlib import Path

st.set_page_config(page_title="Ký tài liệu", page_icon="✍️", layout="wide")
st.title("✍️ Ký tài liệu")

# ─── HẰNG SỐ ──────────────────────────────────────────────────────────────────
_APP_DIR = Path(__file__).parent
CHU_KY_DIR = _APP_DIR / "chu_ky"
CHU_KY_DIR.mkdir(exist_ok=True)
CHU_KY_SAVE_PATH = CHU_KY_DIR / "chu_ky.png"
RENDER_DPI = 130
HA_PRIORITY_MAX = 4   # Từ khóa ưu tiên <= 4 mới được xem là của Hoàng An
PT_TO_PX = RENDER_DPI / 72

TU_KHOA_FILE = _APP_DIR / "tu_khoa.txt"

def _load_keywords() -> list:
    """Đọc tu_khoa.txt. Cột 4 là số ưu tiên (nhỏ = ưu tiên cao hơn)."""
    result = []
    if TU_KHOA_FILE.exists():
        for line in TU_KHOA_FILE.read_text(encoding="utf-8").splitlines():
            line = line.strip()
            if not line or line.startswith("#"):
                continue
            parts = [p.strip() for p in line.split("|")]
            if len(parts) >= 2:
                kw       = parts[0]
                place    = parts[1].lower() if parts[1].lower() in ("below","above") else "below"
                label    = parts[2] if len(parts) >= 3 else kw
                try:
                    priority = int(parts[3]) if len(parts) >= 4 else 5
                except ValueError:
                    priority = 5
                try:
                    v_offset = float(parts[4]) if len(parts) >= 5 else None
                except ValueError:
                    v_offset = None
                result.append({"kw": kw, "label": label, "place": place,
                               "priority": priority, "v_offset": v_offset})
    if not result:
        result = [
            {"kw": "CÔNG TY HOÀNG AN",      "label": "Công ty Hoàng An",       "place": "above", "priority": 1},
            {"kw": "XÁC NHẬN",             "label": "XÁC NHẬN",               "place": "below", "priority": 2},
            {"kw": "Xác nhận đơn đặt hàng","label": "Xác nhận đơn đặt hàng", "place": "below", "priority": 3},
            {"kw": "Xác nhận của bên sản xuất","label": "Xác nhận bên SX",    "place": "below", "priority": 3},
            {"kw": "Xác nhận của nhà sản xuất","label": "Xác nhận nhà SX",    "place": "below", "priority": 3},
            {"kw": "Nguyễn Minh Hoàng",    "label": "Nguyễn Minh Hoàng",     "place": "above", "priority": 4},
            {"kw": "(Ký, họ tên)",          "label": "(Ký, họ tên)",           "place": "below", "priority": 5},
            {"kw": "Tổng Giám đốc",         "label": "Tổng Giám đốc",          "place": "above", "priority": 6},
        ]
    return result

SIGN_KEYWORDS = _load_keywords()

def _strip_accents(s: str) -> str:
    import unicodedata
    return unicodedata.normalize('NFD', s.lower()).encode('ascii', 'ignore').decode('ascii')

# labels không trùng để hiển thị dropdown
_UNIQUE_LABELS = list(dict.fromkeys(e["label"] for e in SIGN_KEYWORDS))
_LABEL_TO_ENTRY = {}
for e in SIGN_KEYWORDS:
    _LABEL_TO_ENTRY.setdefault(e["label"], e)

# ─── SESSION STATE ────────────────────────────────────────────────────────────
for k, v in {
    "sig_active": None, "canvas_key": 0, "show_change_sig": False,
    "batch_results": [], "batch_orig_bytes": {},
}.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _find_chu_ky():
    for pat in ["*.png","*.jpg","*.jpeg","*.PNG","*.JPG"]:
        for f in sorted(CHU_KY_DIR.glob(pat)):
            try: return Image.open(f).convert("RGBA")
            except: continue
    return None

def remove_white_bg(img: Image.Image, threshold=230) -> Image.Image:
    img = img.convert("RGBA"); data = np.array(img)
    m = (data[:,:,0]>threshold)&(data[:,:,1]>threshold)&(data[:,:,2]>threshold)
    data[m,3] = 0
    return Image.fromarray(data,"RGBA")

def save_chu_ky(img: Image.Image):
    img.save(CHU_KY_SAVE_PATH,"PNG"); st.session_state.sig_active = img

# ── Tải file từ đường dẫn (vd ổ Drive đã mount: H:\... hoặc /storage/...) ─────

def load_bytes_from_path(path_str: str):
    """Đọc bytes từ đường dẫn local/đường dẫn ổ Drive đã mount.
    Trả về (bytes, filename) hoặc (None, lỗi)."""
    if not path_str or not path_str.strip():
        return None, None
    p = Path(path_str.strip().strip('"').strip("'"))
    if not p.exists():
        return None, f"Không tìm thấy đường dẫn: `{p}`"
    if p.is_dir():
        return None, f"Đây là thư mục, không phải file: `{p}`"
    try:
        return p.read_bytes(), p.name
    except Exception as e:
        return None, f"Lỗi đọc file: {e}"

# ── Tự động load chữ ký đã lưu trước đó ──────────────────────────────────────
if st.session_state.sig_active is None:
    st.session_state.sig_active = _find_chu_ky()

# ── PDF helpers ───────────────────────────────────────────────────────────────

def render_pdf_page(pdf_bytes: bytes, page_num: int) -> Image.Image:
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    pix = doc[page_num].get_pixmap(matrix=fitz.Matrix(PT_TO_PX,PT_TO_PX), alpha=False)
    return Image.frombytes("RGB",[pix.width,pix.height],pix.samples)

def render_pdf_page_hq(pdf_bytes: bytes, page_num: int, dpi: int = 200) -> Image.Image:
    """Render độ phân giải cao hơn — dùng riêng cho AI đọc chữ rõ hơn."""
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    ratio = dpi / 72
    pix = doc[page_num].get_pixmap(matrix=fitz.Matrix(ratio, ratio), alpha=False)
    return Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

def scan_page_for_keywords(pdf_bytes: bytes, page_num: int) -> list:
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    page = doc[page_num]; found, seen = [], []
    for entry in SIGN_KEYWORDS:
        try: rects = page.search_for(entry["kw"])
        except: continue
        for rect in rects:
            if any(abs(rect.y0-s[0])<8 and abs(rect.x0-s[1])<20 for s in seen): continue
            seen.append((rect.y0,rect.x0))
            found.append({
                "keyword":  entry["kw"],
                "label":    entry["label"],
                "place":    entry["place"],
                "priority": entry.get("priority", 5),
                "v_offset": entry.get("v_offset"),  # None = dùng slider
                "pt": (rect.x0, rect.y0, rect.x1, rect.y1),
                "px": (rect.x0*PT_TO_PX, rect.y0*PT_TO_PX,
                       rect.x1*PT_TO_PX, rect.y1*PT_TO_PX),
            })
    # Ưu tiên nhỏ nhất trước; cùng ưu tiên thì theo Y từ trên xuống
    found.sort(key=lambda a: (a.get("priority", 5), a["pt"][1]))
    return found

def find_content_bottom(page) -> tuple:
    """Trả về (y_bottom, x_center_of_content) — vùng trống ngay dưới nội dung cuối trang."""
    import fitz
    blocks = page.get_text("blocks")
    pw, ph = page.rect.width, page.rect.height
    if not blocks:
        return ph * 0.75, pw / 2
    # Y dưới cùng của nội dung
    y_bottom = max(b[3] for b in blocks)
    # X trung tâm của toàn bộ nội dung (trung bình giữa x0 và x1 của tất cả blocks)
    x_centers = [(b[0] + b[2]) / 2 for b in blocks]
    x_mid = sum(x_centers) / len(x_centers)
    return y_bottom, x_mid


def sign_pdf_bottom_center(pdf_bytes: bytes, sig_img, pages: list,
                           width_pct: float, v_offset: float = 15) -> tuple:
    """Ký ở vùng trống bên dưới nội dung, căn giữa trang — dùng khi không có từ khóa HA."""
    import fitz
    buf = io.BytesIO(); sig_img.save(buf, "PNG"); sig_png = buf.getvalue()
    ar = sig_img.width / sig_img.height
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    signed = []
    for i, page in enumerate(doc):
        if (i + 1) not in pages:
            continue
        pw, ph = page.rect.width, page.rect.height
        sw = pw * width_pct / 100
        sh = sw / ar
        y_bottom, x_mid = find_content_bottom(page)
        x = max(0, min((pw - sw) / 2, pw - sw))   # căn giữa trang
        y = max(0, min(y_bottom + v_offset, ph - sh))
        page.insert_image(fitz.Rect(x, y, x + sw, y + sh), stream=sig_png)
        signed.append(i + 1)
    out = io.BytesIO(); doc.save(out)
    return out.getvalue(), signed


def auto_find_keyword_in_doc(pdf_bytes: bytes):
    """Quét tài liệu, trả về (kw, place) của từ khóa Hoàng An tốt nhất.
    Chỉ chấp nhận keyword có priority <= HA_PRIORITY_MAX.
    Trả về (None, None) nếu không tìm thấy → báo hiệu dùng fallback bottom-center."""
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    candidates = []
    for pg_i, page in enumerate(doc):
        for entry in SIGN_KEYWORDS:
            p = entry.get("priority", 5)
            if p > HA_PRIORITY_MAX:
                continue
            try:
                rects = page.search_for(entry["kw"])
                if rects:
                    candidates.append((p, pg_i, rects[0].y0,
                                       entry["kw"], entry["place"],
                                       entry.get("v_offset")))  # None = dùng default
            except:
                continue
    if not candidates:
        return None, None, None   # → fallback bottom-center
    candidates.sort(key=lambda c: (c[0], c[1], c[2]))
    best = candidates[0]
    return best[3], best[4], best[5]  # kw, place, v_offset

def get_total_pages(pdf_bytes: bytes) -> int:
    import fitz
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    return doc.page_count

def _sig_xy_pt(area, sig_h_pt, v_offset):
    x0,y0,x1,y1 = area["pt"]
    return (x0, y1+v_offset) if area["place"]=="below" else (x0, y0-sig_h_pt-v_offset)

def _sig_xy_px(area, sig_h_px, v_offset_px):
    x0,y0,x1,y1 = area["px"]
    return (x0, y1+v_offset_px) if area["place"]=="below" else (x0, y0-sig_h_px-v_offset_px)

def _manual_xy(bw, bh, sw, sh, position, margin):
    m = margin
    return {"Trên trái":(m,m),"Trên phải":(bw-sw-m,m),
            "Dưới trái":(m,bh-sh-m),"Dưới phải":(bw-sw-m,bh-sh-m),
            "Giữa trang":((bw-sw)//2,(bh-sh)//2),
            "Giữa dưới":((bw-sw)//2,bh-sh-m)}.get(position,(bw-sw-m,bh-sh-m))

# ── Ghép chữ ký ───────────────────────────────────────────────────────────────

def overlay_sig(base: Image.Image, sig: Image.Image,
                x_px, y_px, sw_px, highlight_px=None) -> Image.Image:
    bw,bh = base.size
    sw = max(30, min(sw_px, bw)); sh = int(sig.height*sw/sig.width)
    sig_r = sig.resize((sw,sh),Image.LANCZOS).convert("RGBA")
    px = max(0,min(int(x_px),bw-sw)); py = max(0,min(int(y_px),bh-sh))
    result = base.convert("RGBA")
    if highlight_px:
        hl = Image.new("RGBA",result.size,(0,0,0,0)); draw = ImageDraw.Draw(hl)
        x0,y0,x1,y1 = [int(v) for v in highlight_px]
        draw.rectangle([x0,y0,x1,y1],outline=(255,140,0,230),width=3)
        result = Image.alpha_composite(result,hl)
    result.paste(sig_r,(px,py),mask=sig_r)
    return result.convert("RGB")

def sign_pdf_auto(pdf_bytes, sig_img, pages, keyword, place, width_pct, v_offset):
    import fitz
    buf=io.BytesIO(); sig_img.save(buf,"PNG"); sig_png=buf.getvalue()
    ar=sig_img.width/sig_img.height
    doc=fitz.open(stream=pdf_bytes,filetype="pdf")
    signed,skipped=[],[]
    for i,page in enumerate(doc):
        pg=i+1
        if pg not in pages: continue
        pw,ph=page.rect.width,page.rect.height
        sw=pw*width_pct/100; sh=sw/ar
        rects=page.search_for(keyword)
        if not rects: skipped.append(pg); continue
        r=rects[0]
        area={"place":place,"pt":(r.x0,r.y0,r.x1,r.y1)}
        x,y=_sig_xy_pt(area,sh,float(v_offset))
        x=max(0,min(x,pw-sw)); y=max(0,min(y,ph-sh))
        page.insert_image(fitz.Rect(x,y,x+sw,y+sh),stream=sig_png)
        signed.append(pg)
    out=io.BytesIO(); doc.save(out)
    return out.getvalue(), signed, skipped

def sign_image_file(img_bytes: bytes, img_name: str, sig_img,
                    position: str, width_pct: float, margin: int) -> bytes:
    """Ghép chữ ký lên file ảnh, trả về PNG bytes."""
    base = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    bw, bh = base.size
    sw = max(30, int(bw * width_pct / 100))
    sh = int(sig_img.height * sw / sig_img.width)
    x, y = _manual_xy(bw, bh, sw, sh, position, margin)
    result = overlay_sig(base, sig_img, x, y, sw)
    buf = io.BytesIO()
    result.save(buf, "PNG")
    return buf.getvalue()

def scan_image_for_keywords(img: Image.Image) -> list:
    """OCR ảnh để tìm vị trí từ khóa. Trả về cùng format với scan_page_for_keywords."""
    try:
        import pytesseract
    except ImportError:
        return []
    # Chỉ đường dẫn Tesseract + tessdata tiếng Việt nằm trong thư mục app
    import sys, os
    if sys.platform == "win32":
        _tess = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
        if os.path.exists(_tess):
            pytesseract.pytesseract.tesseract_cmd = _tess
    _local_tessdata = str(_APP_DIR / "tessdata")
    if os.path.isdir(_local_tessdata):
        os.environ["TESSDATA_PREFIX"] = _local_tessdata
        lang_str = 'vie+eng'
    else:
        try:
            langs = pytesseract.get_languages()
            lang_str = 'vie+eng' if 'vie' in langs else 'eng'
        except Exception:
            lang_str = 'eng'
    try:
        data = pytesseract.image_to_data(img, output_type=pytesseract.Output.DICT,
                                          lang=lang_str, config='--psm 6')
    except Exception:
        return []
    from collections import defaultdict
    lines = defaultdict(lambda: {'texts': [], 'boxes': []})
    for i in range(len(data['text'])):
        if int(data['conf'][i]) < 0:
            continue
        key = (data['page_num'][i], data['block_num'][i],
               data['par_num'][i],  data['line_num'][i])
        t = data['text'][i].strip()
        if t:
            lines[key]['texts'].append(t)
            lines[key]['boxes'].append((
                data['left'][i], data['top'][i],
                data['left'][i] + data['width'][i],
                data['top'][i] + data['height'][i],
            ))
    found, seen = [], []
    for key in sorted(lines.keys()):
        ld = lines[key]
        if not ld['texts']:
            continue
        line_norm = _strip_accents(' '.join(ld['texts']))
        for entry in SIGN_KEYWORDS:
            if _strip_accents(entry['kw']) in line_norm:
                x0 = min(b[0] for b in ld['boxes'])
                y0 = min(b[1] for b in ld['boxes'])
                x1 = max(b[2] for b in ld['boxes'])
                y1 = max(b[3] for b in ld['boxes'])
                if any(abs(y0 - s) < 8 for s in seen):
                    continue
                seen.append(y0)
                found.append({
                    'keyword': entry['kw'], 'label': entry['label'],
                    'place': entry['place'], 'priority': entry.get('priority', 5),
                    'v_offset': entry.get('v_offset'),
                    'px': (x0, y0, x1, y1), 'pt': (x0, y0, x1, y1),
                })
                break
    found.sort(key=lambda a: (a.get('priority', 5), a['px'][1]))
    return found

# ── AI Vision (Claude API) — tìm vị trí ký khi local scan không ra ──────────

def _get_anthropic_key():
    try:
        return st.secrets.get("ANTHROPIC_API_KEY", "")
    except Exception:
        return ""

def ai_find_signature_position(img: Image.Image, sig_img: Image.Image, width_pct: float = 22):
    """Gọi Claude API (vision) để tìm KHUNG dòng chữ chứa từ khóa cần ký bên cạnh,
    sau đó TỰ TÍNH vị trí đặt chữ ký (không để AI đoán trực tiếp vị trí — kém chính xác hơn).
    Trả về (dict{x_pct,y_pct,reason}, None) khi thành công, hoặc (None, thông_báo_lỗi)."""
    api_key = _get_anthropic_key()
    if not api_key:
        return None, ("Chưa cấu hình ANTHROPIC_API_KEY. Vào Streamlit Cloud → App → "
                       "Settings → Secrets, thêm dòng:\nANTHROPIC_API_KEY = \"sk-ant-...\"")

    import base64, json, requests

    img_s = img.convert("RGB")
    max_dim = 2000
    if max(img_s.size) > max_dim:
        ratio = max_dim / max(img_s.size)
        img_s = img_s.resize((int(img_s.width * ratio), int(img_s.height * ratio)))
    buf = io.BytesIO(); img_s.save(buf, "JPEG", quality=92)
    b64 = base64.b64encode(buf.getvalue()).decode()

    kw_hint = ", ".join(f'"{e["kw"]}"' for e in SIGN_KEYWORDS)
    extra_kw = '"Xác nhận của bên sản xuất", "Xác nhận của nhà sản xuất", "Người mua hàng", "Ký, ghi rõ họ tên", "Đại diện bên"'
    prompt = (
        "Đây là ảnh 1 trang tài liệu/đơn hàng tiếng Việt cần ký tên. "
        f"Tìm DÒNG CHỮ phù hợp nhất để đặt chữ ký ngay sát nó.\n\n"
        f"DANH SÁCH TỪ KHÓA ƯU TIÊN (theo thứ tự ưu tiên giảm dần):\n{kw_hint}\n\n"
        f"TỪ KHÓA BỔ SUNG (cũng tìm nếu không thấy từ khóa trên):\n{extra_kw}\n\n"
        "QUY TẮC:\n"
        "1. Ưu tiên tìm CHÍNH XÁC các từ khóa trên trong trang\n"
        "2. Nếu tìm thấy nhiều từ khóa, chọn từ khóa có THỨ TỰ ƯU TIÊN CAO NHẤT (đầu danh sách)\n"
        "3. Nếu không tìm thấy từ khóa nào, chọn dòng chữ hợp lý nhất gần cuối trang "
        "(ví dụ dòng tên người/chức danh/ô trống cuối văn bản)\n\n"
        "Trả lời CHỈ một JSON object, không giải thích, không markdown, đúng định dạng:\n"
        '{"found": true, "matched_text": "<nguyên văn dòng chữ đã chọn>", '
        '"box_pct": [x0, y0, x1, y1], "place": "below"}\n'
        "Trong đó box_pct là tọa độ % (0-100) của khung SÁT QUANH dòng chữ đó so với chiều "
        "rộng/cao toàn ảnh (x0,y0 = góc trên-trái; x1,y1 = góc dưới-phải; x0<x1; y0<y1). "
        "\"place\" là \"below\" nếu nên ký ngay DƯỚI dòng này, hoặc \"above\" nếu nên ký ngay TRÊN.\n"
        'Nếu trang không có dòng chữ nào hợp lý, trả {"found": false, "reason": "..."}'
    )

    try:
        resp = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "x-api-key": api_key,
                "anthropic-version": "2023-06-01",
                "content-type": "application/json",
            },
            json={
                "model": "claude-sonnet-4-6",
                "max_tokens": 400,
                "messages": [{
                    "role": "user",
                    "content": [
                        {"type": "image", "source": {"type": "base64",
                                                      "media_type": "image/jpeg", "data": b64}},
                        {"type": "text", "text": prompt},
                    ],
                }],
            },
            timeout=40,
        )
        if resp.status_code != 200:
            return None, f"Lỗi API ({resp.status_code}): {resp.text[:200]}"
        data = resp.json()
        text = "".join(b.get("text", "") for b in data.get("content", []) if b.get("type") == "text")
        text = text.strip().replace("```json", "").replace("```", "").strip()
        parsed = json.loads(text)
        if not parsed.get("found"):
            return None, f"AI không tìm thấy dòng chữ phù hợp ({parsed.get('reason','')})."

        box = parsed.get("box_pct")
        if not box or len(box) != 4:
            return None, "AI trả về box_pct không hợp lệ."
        x0, y0, x1, y1 = [max(0, min(100, float(v))) for v in box]
        place = parsed.get("place", "below")
        if place not in ("below", "above"):
            place = "below"

        # Tự tính vị trí đặt chữ ký dựa trên khung dòng chữ (chính xác hơn để AI tự đoán)
        page_ar = img.width / img.height           # tỉ lệ trang (cùng tỉ lệ với ảnh resize)
        sig_ar  = sig_img.width / sig_img.height
        gap_pct = 1.2                                # khoảng cách nhỏ với dòng chữ, % chiều cao trang
        sig_h_pct = width_pct * page_ar / sig_ar     # chiều cao chữ ký quy theo % chiều cao trang

        if place == "below":
            x_pct, y_pct = x0, y1 + gap_pct
        else:
            x_pct, y_pct = x0, y0 - sig_h_pct - gap_pct

        x_pct = max(0, min(95, x_pct))
        y_pct = max(0, min(95, y_pct))
        reason = parsed.get("matched_text", "")
        return {"x_pct": x_pct, "y_pct": y_pct, "reason": reason}, None
    except json.JSONDecodeError:
        return None, "AI trả về định dạng không đọc được. Thử lại."
    except Exception as e:
        return None, f"Lỗi gọi AI: {e}"

def sign_pdf_at_xy_pct(pdf_bytes: bytes, sig_img, pages: list,
                       x_pct: float, y_pct: float, width_pct: float) -> bytes:
    """Ký PDF tại vị trí % tùy ý (dùng cho kết quả AI hoặc kéo thả tự do)."""
    import fitz
    buf = io.BytesIO(); sig_img.save(buf, "PNG"); sig_png = buf.getvalue()
    ar = sig_img.width / sig_img.height
    doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    for i, page in enumerate(doc):
        if (i + 1) not in pages:
            continue
        pw, ph = page.rect.width, page.rect.height
        sw = pw * width_pct / 100; sh = sw / ar
        x = pw * x_pct / 100; y = ph * y_pct / 100
        x = max(0, min(x, pw - sw)); y = max(0, min(y, ph - sh))
        page.insert_image(fitz.Rect(x, y, x + sw, y + sh), stream=sig_png)
    out = io.BytesIO(); doc.save(out)
    return out.getvalue()

def sign_image_auto(img_bytes: bytes, sig_img, area: dict,
                    width_pct: float, v_offset: float) -> bytes:
    """Ký ảnh theo vị trí từ khóa OCR."""
    base = Image.open(io.BytesIO(img_bytes)).convert("RGB")
    bw, bh = base.size
    sw = max(30, int(bw * width_pct / 100))
    sh = int(sig_img.height * sw / sig_img.width)
    eff_v = area['v_offset'] if area.get('v_offset') is not None else v_offset
    x_px, y_px = _sig_xy_px(area, sh, eff_v)
    result = overlay_sig(base, sig_img, x_px, y_px, sw)
    buf = io.BytesIO()
    result.save(buf, "PNG")
    return buf.getvalue()

_SIGN_KW_XL = [
    "công ty hoàng an", "cong ty hoang an",
    "nguyễn minh hoàng", "nguyen minh hoang",
    "bên sản xuất", "nhà sản xuất", "xác nhận nhà sản xuất",
]

def find_excel_sig_position(excel_bytes: bytes) -> tuple:
    """Tìm row/col/colOff để ký ngay dưới keyword, căn giữa vùng merge.
    Trả về (row_1based, col_0based, found_text, col_off_emu)."""
    from openpyxl import load_workbook
    try:
        # load_workbook không save → không corrupt file
        wb = load_workbook(io.BytesIO(excel_bytes), data_only=True)
        ws = wb.active

        target_row = target_col = None
        target_text = ""
        for row_cells in ws.iter_rows():
            for cell in row_cells:
                if not (cell.value and isinstance(cell.value, str)):
                    continue
                val = cell.value.strip()
                for kw in _SIGN_KW_XL:
                    if kw in val.lower():
                        target_row  = cell.row      # 1-based
                        target_col  = cell.column   # 1-based
                        target_text = val[:50]
                        break
                if target_row:
                    break
            if target_row:
                break

        if not target_row:
            wb.close()
            return None, None, None, 0

        # Tìm vùng merge chứa cell này
        min_col, max_col = target_col, target_col
        for merge in ws.merged_cells.ranges:
            if (merge.min_row <= target_row <= merge.max_row and
                    merge.min_col <= target_col <= merge.max_col):
                min_col = merge.min_col
                max_col = merge.max_col
                break

        # Tính tổng chiều rộng vùng merge → đặt chữ ký căn giữa
        def _col_emu(col_idx):
            from openpyxl.utils import get_column_letter
            letter = get_column_letter(col_idx)
            dim = ws.column_dimensions.get(letter)
            px = int((dim.width if dim and dim.width else 8.43) * 7 + 5)
            return px * 9525

        total_w_emu = sum(_col_emu(c) for c in range(min_col, max_col + 1))
        sig_w_emu   = 160 * 9525
        col_off_emu = max(0, (total_w_emu - sig_w_emu) // 2)

        wb.close()
        # row 1-based (dòng ký = ngay dưới keyword), col 0-based
        return target_row, min_col - 1, target_text, col_off_emu

    except Exception:
        pass
    return None, None, None, 0


def sign_excel_file(excel_bytes: bytes, sig_img, position: str,
                    width_px: int = 160, height_px: int = 80) -> tuple:
    """
    Chèn chữ ký vào xlsx bằng zipfile trực tiếp.
    - Tự tìm vị trí keyword (CÔNG TY HOÀNG AN, v.v.)
    - oneCellAnchor + giữ đúng tỉ lệ ảnh → không méo
    - Không dùng openpyxl save → không corrupt drawing cũ
    Trả về (bytes, info_str).
    """
    import zipfile, re
    from openpyxl import load_workbook

    # ── Tìm vị trí ký trong file ──────────────────────────────────────────────
    sig_row_0, sig_col_0, found_text, col_off_emu = find_excel_sig_position(excel_bytes)

    if sig_row_0 is None:
        # Fallback: dùng position tham số
        try:
            wb_ro = load_workbook(io.BytesIO(excel_bytes), read_only=True)
            max_row = max((ws.max_row or 30 for ws in wb_ro.worksheets), default=30)
            wb_ro.close()
        except Exception:
            max_row = 40
        col_map = {"Dưới phải":7,"Trên phải":7,"Dưới trái":0,"Trên trái":0,
                   "Giữa trang":3,"Giữa dưới":3}
        row_map = {"Dưới phải":max_row,"Dưới trái":max_row,"Giữa dưới":max_row,
                   "Trên phải":0,"Trên trái":0,"Giữa trang":max(0,max_row//2)}
        sig_col_0   = col_map.get(position, 7)
        sig_row_0   = row_map.get(position, max_row)
        col_off_emu = 0
        info = f"Không tìm thấy keyword → đặt theo vị trí {position}"
    else:
        info = f"Tìm thấy [{found_text}] → ký row {sig_row_0+1}, căn giữa (offset {col_off_emu//9525}px)"

    # ── Chuẩn bị ảnh, giữ đúng tỉ lệ (oneCellAnchor) ─────────────────────────
    sig_buf = io.BytesIO()
    sig_img.convert("RGBA").save(sig_buf, "PNG")
    sig_png = sig_buf.getvalue()
    # EMU: dùng width cố định, tính height theo tỉ lệ ảnh
    w_emu = width_px * 9525
    h_emu = int(sig_img.height / sig_img.width * w_emu)

    in_buf = io.BytesIO(excel_bytes)
    out_buf = io.BytesIO()

    with zipfile.ZipFile(in_buf, 'r') as zin:
        all_names = zin.namelist()

        # Tên ảnh mới tránh trùng
        existing_imgs = [n for n in all_names if n.startswith('xl/media/image')]
        new_idx = len(existing_imgs) + 1
        new_img_path = f"xl/media/image{new_idx}.png"
        while new_img_path in all_names:
            new_idx += 1; new_img_path = f"xl/media/image{new_idx}.png"
        new_img_base = new_img_path.split('/')[-1]

        drw_path = next(
            (n for n in all_names
             if n.startswith('xl/drawings/drawing') and n.endswith('.xml')
             and '_rels' not in n), None)
        if drw_path is None:
            return excel_bytes, "Không có drawing trong file"

        drw_rels_path = f"xl/drawings/_rels/{drw_path.split('/')[-1]}.rels"
        drw_xml  = zin.read(drw_path).decode('utf-8')
        drw_rels = (zin.read(drw_rels_path).decode('utf-8') if drw_rels_path in all_names
                    else '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>'
                         '<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships">'
                         '</Relationships>')

        used_rids  = [int(x) for x in re.findall(r'Id="rId(\d+)"', drw_rels)]
        new_rid    = f"rId{max(used_rids, default=0) + 1}"
        used_ids   = [int(x) for x in re.findall(r'\bid="(\d+)"', drw_xml)]
        new_pic_id = max(used_ids, default=1) + 1

        # oneCellAnchor: ảnh kích thước cố định, không bị kéo dãn
        anchor = (
            f'<xdr:oneCellAnchor>'
            f'<xdr:from>'
            f'<xdr:col>{sig_col_0}</xdr:col><xdr:colOff>{col_off_emu}</xdr:colOff>'
            f'<xdr:row>{sig_row_0}</xdr:row><xdr:rowOff>57150</xdr:rowOff>'
            f'</xdr:from>'
            f'<xdr:ext cx="{w_emu}" cy="{h_emu}"/>'
            f'<xdr:pic><xdr:nvPicPr>'
            f'<xdr:cNvPr id="{new_pic_id}" name="ChuKyHoangAn"/>'
            f'<xdr:cNvPicPr><a:picLocks noChangeAspect="1"/></xdr:cNvPicPr>'
            f'</xdr:nvPicPr><xdr:blipFill>'
            f'<a:blip xmlns:r="http://schemas.openxmlformats.org/officeDocument/2006/relationships"'
            f' r:embed="{new_rid}"/>'
            f'<a:stretch><a:fillRect/></a:stretch>'
            f'</xdr:blipFill><xdr:spPr>'
            f'<a:xfrm><a:off x="0" y="0"/><a:ext cx="{w_emu}" cy="{h_emu}"/></a:xfrm>'
            f'<a:prstGeom prst="rect"><a:avLst/></a:prstGeom>'
            f'</xdr:spPr></xdr:pic><xdr:clientData/>'
            f'</xdr:oneCellAnchor>'
        )
        new_drw_xml  = drw_xml.replace('</xdr:wsDr>', anchor + '</xdr:wsDr>')
        new_rel      = (f'<Relationship Id="{new_rid}"'
                        f' Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/image"'
                        f' Target="../media/{new_img_base}"/>')
        new_drw_rels = drw_rels.replace('</Relationships>', new_rel + '</Relationships>')

        with zipfile.ZipFile(out_buf, 'w', zipfile.ZIP_DEFLATED) as zout:
            for name in all_names:
                if name == drw_path:
                    zout.writestr(name, new_drw_xml.encode('utf-8'))
                elif name == drw_rels_path:
                    zout.writestr(name, new_drw_rels.encode('utf-8'))
                else:
                    zout.writestr(name, zin.read(name))
            zout.writestr(new_img_path, sig_png)

    return out_buf.getvalue(), info

def sign_word_file(docx_bytes: bytes, sig_img, width_cm: float = 4.0) -> tuple:
    """Chèn chữ ký vào Word (.docx). Tìm keyword rồi thêm ảnh sau đó, hoặc cuối tài liệu."""
    from docx import Document
    from docx.shared import Cm
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph as DocxParagraph

    doc = Document(io.BytesIO(docx_bytes))

    sig_buf = io.BytesIO()
    sig_img.convert("RGB").save(sig_buf, "PNG")
    sig_buf.seek(0)

    target_para = None
    found_kw = None

    def _search(paras):
        nonlocal target_para, found_kw
        for para in paras:
            t = para.text.lower()
            for kw in _SIGN_KW_XL:
                if kw in t:
                    target_para = para; found_kw = kw; return True
        return False

    _search(doc.paragraphs)
    if not target_para:
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if _search(cell.paragraphs): break
                if target_para: break
            if target_para: break

    if target_para:
        new_p = OxmlElement('w:p')
        target_para._p.addnext(new_p)
        new_para = DocxParagraph(new_p, target_para._parent)
        run = new_para.add_run()
        run.add_picture(sig_buf, width=Cm(width_cm))
        info = f"Tìm thấy [{found_kw}] → ký ngay dưới"
    else:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(sig_buf, width=Cm(width_cm))
        info = "Không tìm thấy keyword → ký ở cuối tài liệu"

    out_buf = io.BytesIO()
    doc.save(out_buf)
    return out_buf.getvalue(), info

def create_zip(results: list) -> bytes:
    """Tạo ZIP từ danh sách kết quả ký hàng loạt."""
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        for r in results:
            if r.get("bytes"):
                zf.writestr(r.get("out_name", r["name"]), r["bytes"])
    return buf.getvalue()

def _ext(name: str) -> str:
    return Path(name).suffix.lower()

def process_one_file(name: str, data: bytes, sig_img, kw, place,
                     pages_mode: str, width_pct: float, v_offset: float,
                     kw_v_offset=None,
                     img_position: str = "Dưới phải", img_margin: int = 35) -> dict:
    """Ký một file (PDF / ảnh / Excel). Trả về dict kết quả."""
    ext = _ext(name)
    try:
        if ext == ".pdf":
            total = get_total_pages(data)
            pages = list(range(1, total+1)) if pages_mode == "Tất cả trang" else [total]
            if kw is None:
                out, signed = sign_pdf_bottom_center(data, sig_img, pages, width_pct)
                detail = f"Ký {len(signed)} trang (vùng trống cuối — không có từ khóa HA)"
            else:
                # Ưu tiên v_offset từ keyword (tu_khoa.txt), sau đó mới dùng slider
                effective_offset = kw_v_offset if kw_v_offset is not None else v_offset
                out, signed, skipped = sign_pdf_auto(data, sig_img, pages, kw, place,
                                                      width_pct, effective_offset)
                offset_note = f" [kc={int(effective_offset)}pt]" if kw_v_offset is not None else ""
                detail = f"Ký {len(signed)} trang theo [{kw}]{offset_note}"
                if skipped:
                    detail += f" | bỏ qua trang {skipped}"
            out_name = name
        elif ext in (".png", ".jpg", ".jpeg"):
            _base_img = Image.open(io.BytesIO(data)).convert("RGB")
            _img_areas = scan_image_for_keywords(_base_img)
            _ha = [a for a in _img_areas if a.get("priority", 5) <= HA_PRIORITY_MAX]
            if _ha:
                out = sign_image_auto(data, sig_img, _ha[0], width_pct, v_offset)
                detail = f"OCR: [{_ha[0]['keyword']}] → ký bên dưới"
            else:
                out = sign_image_file(data, name, sig_img, img_position, width_pct, img_margin)
                detail = "Ghép chữ ký (vị trí thủ công)"
            out_name = f"{Path(name).stem}.png"
        elif ext in (".xlsx", ".xls"):
            out, xl_info = sign_excel_file(data, sig_img, img_position)
            detail = xl_info
            out_name = name
        elif ext == ".docx":
            out, word_info = sign_word_file(data, sig_img)
            detail = word_info
            out_name = name
        else:
            return {"name": name, "status": "⏭️ Bỏ qua",
                    "detail": f"Định dạng {ext} chưa hỗ trợ", "bytes": None, "out_name": name}

        return {"name": name, "status": "✅ OK", "detail": detail,
                "bytes": out, "out_name": out_name}
    except Exception as e:
        return {"name": name, "status": "❌ Lỗi", "detail": str(e),
                "bytes": None, "out_name": name}

# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR — Chữ ký
# ═══════════════════════════════════════════════════════════════════════════════
with st.sidebar:
    st.header("✍️ Chữ ký")
    if st.session_state.sig_active and not st.session_state.show_change_sig:
        st.success("Đang dùng chữ ký đã lưu")
        st.image(st.session_state.sig_active, use_container_width=True)
        if st.button("🔄 Đổi chữ ký", use_container_width=True):
            st.session_state.show_change_sig = True; st.rerun()
    else:
        if st.session_state.show_change_sig: st.info("Chọn chữ ký mới:")
        tab_up, tab_draw = st.tabs(["📤 Upload ảnh","🖊 Vẽ tay"])
        new_sig = None
        with tab_up:
            sf = st.file_uploader("Ảnh chữ ký (PNG/JPG)", type=["png","jpg","jpeg"])
            sig_path = st.text_input(
                "...hoặc dán đường dẫn file (ổ Drive đã mount, vd H:\\...\\chu_ky.png)",
                key="sig_path_input")
            loaded = None
            if sf:
                loaded = Image.open(sf)
            elif sig_path:
                _b, _name_or_err = load_bytes_from_path(sig_path)
                if _b:
                    loaded = Image.open(io.BytesIO(_b))
                elif _name_or_err:
                    st.error(_name_or_err)
            if loaded:
                rm = st.checkbox("Xóa nền trắng", value=True)
                new_sig = remove_white_bg(loaded) if rm else loaded.convert("RGBA")
                st.image(new_sig, use_container_width=True)
        with tab_draw:
            try:
                from streamlit_drawable_canvas import st_canvas
                c1,c2=st.columns([4,1])
                with c2:
                    if st.button("Xóa",use_container_width=True):
                        st.session_state.canvas_key+=1; st.rerun()
                cr=st_canvas(stroke_width=3,stroke_color="#111111",
                             background_color="#f5f5f5",height=150,
                             drawing_mode="freedraw",
                             key=f"cv_{st.session_state.canvas_key}",
                             display_toolbar=False)
                if cr.image_data is not None:
                    arr=cr.image_data.astype(np.uint8)
                    if (arr[:,:,:3].min(axis=2)<180).any():
                        new_sig=remove_white_bg(Image.fromarray(arr[:,:,:3]),200)
            except ImportError:
                st.info("Tính năng vẽ tay không khả dụng.")
        if new_sig is not None:
            st.divider()
            if st.button("💾 Lưu & dùng chữ ký này",type="primary",use_container_width=True):
                save_chu_ky(new_sig); st.session_state.show_change_sig=False; st.rerun()
        if st.session_state.show_change_sig:
            if st.button("↩️ Giữ chữ ký cũ",use_container_width=True):
                st.session_state.show_change_sig=False; st.rerun()

# ═══════════════════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════════════════
if st.session_state.sig_active is None:
    st.info("Chưa có chữ ký — upload hoặc vẽ chữ ký ở thanh bên trái rồi lưu.")
    st.stop()

# ═══════════════════════════════════════════════════════════════════════════════
# TAB 2 — Ký hàng loạt
# ═══════════════════════════════════════════════════════════════════════════════
col_b1, col_b2 = st.columns([1,1], gap="large")

with col_b1:
    # ── Chọn file ────────────────────────────────────────────────
    uploaded_many = st.file_uploader(
        "📂 Chọn file cần ký (PDF, ảnh, Excel, Word)",
        type=["pdf","png","jpg","jpeg","xlsx","xls","docx"],
        accept_multiple_files=True, key="b_upload")

    files_data = []
    if uploaded_many:
        files_data = [(f.name, f.read()) for f in uploaded_many]

    if files_data:
        st.success(f"✅ {len(files_data)} file sẵn sàng")
        with st.expander(f"Danh sách {len(files_data)} file"):
            for name, data in files_data:
                st.text(f"  • {name}  ({len(data)//1024} KB)")

    # ── Cài đặt (mặc định đóng) ──────────────────────────────────────
    with st.expander("⚙️ Cài đặt ký", expanded=False):
        kw_mode = st.radio("Từ khóa tìm vị trí",
                          ["🤖 Tự động", "🎯 Chọn cụ thể"],
                          horizontal=True, key="b_kwmode")

        specific_kw, specific_place = None, None
        if kw_mode == "🎯 Chọn cụ thể":
            chosen_label = st.selectbox("Từ khóa", _UNIQUE_LABELS, key="b_kwsel")
            entry = _LABEL_TO_ENTRY[chosen_label]
            specific_kw, specific_place = entry["kw"], entry["place"]

        pages_mode = st.radio("Trang ký PDF",["Tất cả trang","Chỉ trang cuối"],
                             horizontal=True, key="b_pgmode")

        _cw1, _cw2 = st.columns(2)
        with _cw1:
            b_width = st.slider("Kích thước (%rộng)",5,60,22,key="b_w")
        with _cw2:
            b_voffset = st.slider("Khoảng cách (pt)",-30,80,4,key="b_vo")

        _cw3, _cw4 = st.columns(2)
        with _cw3:
            b_img_pos = st.selectbox("Vị trí ảnh/Excel",
                ["Dưới phải","Dưới trái","Giữa dưới","Trên phải","Trên trái"],key="b_ipos")
        with _cw4:
            b_margin = st.slider("Lề ảnh (px)",5,150,35,key="b_mg")

    # ── Bắt đầu ký ───────────────────────────────────────────────────
    ready = bool(files_data)
    if not ready:
        st.caption("Upload file để bắt đầu.")

    if ready and st.button("▶️ Bắt đầu ký hàng loạt",type="primary",
                           use_container_width=True, key="b_go"):
        sig = st.session_state.sig_active
        results = []
        progress_bar = st.progress(0, text="Đang xử lý...")
        status_box = st.empty()

        for i, (name, data) in enumerate(files_data):
            status_box.text(f"Đang xử lý: {name}  ({i+1}/{len(files_data)})")

            # Tìm keyword
            if kw_mode == "🤖 Tự động":
                kw, place, kw_voff = auto_find_keyword_in_doc(data)
            else:
                kw, place, kw_voff = specific_kw, specific_place, None

            r = process_one_file(name, data, sig, kw, place,
                                 pages_mode, b_width, b_voffset,
                                 kw_v_offset=kw_voff,
                                 img_position=b_img_pos, img_margin=b_margin)
            results.append(r)
            progress_bar.progress((i+1)/len(files_data),
                                  text=f"Xong {i+1}/{len(files_data)}: {name}")

        st.session_state.batch_results = results
        st.session_state.batch_orig_bytes = {name: data for name, data in files_data}
        status_box.empty()
        progress_bar.progress(1.0, text="Hoàn tất!")

with col_b2:
    st.subheader("📊 Kết quả & xem trước")
    results = st.session_state.batch_results

    if not results:
        st.info("Kết quả sẽ hiển thị ở đây sau khi ký xong.")
    else:
        ok   = [r for r in results if r["status"] == "✅ OK"]
        skip = [r for r in results if "Bỏ qua" in r["status"]]
        err  = [r for r in results if r["status"] == "❌ Lỗi"]

        c1, c2, c3 = st.columns(3)
        c1.metric("✅ Thành công", len(ok))
        c2.metric("⏭️ Bỏ qua",    len(skip))
        c3.metric("❌ Lỗi",        len(err))

        import pandas as pd
        df = pd.DataFrame([{"File": r["name"], "Trạng thái": r["status"],
                             "Chi tiết": r["detail"]} for r in results])
        st.dataframe(df, use_container_width=True, hide_index=True)

        # ── Xem trước kết quả ────────────────────────────────────────────
        if ok:
            st.divider()
            st.markdown("**👁 Xem trước file đã ký**")
            prev_names = [r["name"] for r in ok]
            prev_sel = st.selectbox("Chọn file để xem trước",
                                    prev_names, key="b_prev_sel")
            sel_r = next(r for r in ok if r["name"] == prev_sel)
            ext_p = _ext(sel_r["name"])

            if ext_p == ".pdf":
                try:
                    total_p = get_total_pages(sel_r["bytes"])
                    if total_p > 1:
                        pg_sel = st.select_slider(
                            "Trang", options=list(range(1, total_p+1)),
                            value=total_p, key="b_pg_prev")
                    else:
                        pg_sel = 1
                    prev_img = render_pdf_page(sel_r["bytes"], pg_sel - 1)
                    st.image(prev_img, use_container_width=True,
                             caption=f"{sel_r['name']} — Trang {pg_sel}/{total_p}")
                except Exception as e:
                    st.error(f"Lỗi xem trước: {e}")

            elif ext_p in (".png", ".jpg", ".jpeg"):
                st.image(sel_r["bytes"], use_container_width=True,
                         caption=sel_r["name"])

            else:
                # Excel: hiện bảng dữ liệu + vị trí chữ ký
                try:
                    import pandas as pd
                    from openpyxl import load_workbook as _lw
                    wb_p = _lw(io.BytesIO(sel_r["bytes"]), read_only=True, data_only=True)
                    ws_p = wb_p.active
                    rows_p = [[c.value for c in r] for r in ws_p.iter_rows(max_row=min(ws_p.max_row or 30, 50))]
                    wb_p.close()
                    df_p = pd.DataFrame(rows_p).fillna("")
                    st.dataframe(df_p, use_container_width=True, hide_index=True)
                    # Tìm vị trí chữ ký
                    row_s, col_s, txt_s, off_s = find_excel_sig_position(sel_r["bytes"])
                    if row_s:
                        st.success(f"Chữ ký đặt tại **Row {row_s+1}**, căn giữa bên dưới \"{txt_s}\" (offset {off_s//9525}px)")
                    else:
                        st.info("Không tìm thấy keyword — chữ ký đặt cuối trang.")
                except Exception as ep:
                    st.info(f"Không xem trước được: {ep}")

            # ── Chỉnh vị trí cho file đang xem ──────────────────────────
            _orig_b = st.session_state.batch_orig_bytes.get(sel_r["name"])
            if _orig_b and ext_p in (".pdf", ".png", ".jpg", ".jpeg"):
                with st.expander("✏️ Chỉnh vị trí chữ ký cho file này"):
                    _sig_e = st.session_state.sig_active

                    _key_ok = bool(_get_anthropic_key())
                    st.caption("🔑 AI Vision: " + ("✅ Đã cấu hình API key" if _key_ok
                              else "❌ Chưa thấy API key (kiểm tra Secrets + Reboot app)"))

                    # ── 🤖 AI tìm vị trí (chạy TRƯỚC khi tạo slider để set session_state an toàn)
                    if ext_p == ".pdf":
                        _epg_preview = st.session_state.get("b_epg", get_total_pages(_orig_b))
                    if st.button("🤖 Dùng AI tìm vị trí", use_container_width=True, key="b_ai_find"):
                        with st.spinner("Đang hỏi AI..."):
                            try:
                                if ext_p == ".pdf":
                                    _ai_img = render_pdf_page_hq(_orig_b, _epg_preview - 1)
                                else:
                                    _ai_img = Image.open(io.BytesIO(_orig_b)).convert("RGB")
                                _ai_width = st.session_state.get("b_ew", 22)
                                _ai_res, _ai_err = ai_find_signature_position(
                                    _ai_img, _sig_e, _ai_width)
                            except Exception as _ai_ex:
                                _ai_res, _ai_err = None, str(_ai_ex)
                        if _ai_res:
                            st.session_state["b_ex"] = round(_ai_res["x_pct"])
                            st.session_state["b_ey"] = round(_ai_res["y_pct"])
                            st.session_state["_b_ai_reason"] = _ai_res.get("reason", "")
                        elif _ai_err:
                            st.session_state["_b_ai_reason"] = None
                            st.error(_ai_err)
                    if st.session_state.get("_b_ai_reason"):
                        st.success(f"🤖 AI đã định vị: {st.session_state['_b_ai_reason']}")

                    _ce1, _ce2 = st.columns(2)
                    with _ce1:
                        _ew = st.slider("Kích thước (% chiều rộng)", 5, 60, 22, key="b_ew")
                    with _ce2:
                        if ext_p == ".pdf":
                            _etotal = get_total_pages(_orig_b)
                            _epg = st.selectbox("Trang", list(range(1, _etotal+1)),
                                                index=_etotal-1, key="b_epg")
                    _ex = st.slider("↔ Ngang (% từ trái)", 0, 95, 65, key="b_ex")
                    _ey = st.slider("↕ Dọc (% từ trên)",   0, 95, 70, key="b_ey")

                    # Preview cập nhật realtime khi kéo slider
                    try:
                        if ext_p == ".pdf":
                            _base_e = render_pdf_page(_orig_b, _epg - 1)
                        else:
                            _base_e = Image.open(io.BytesIO(_orig_b)).convert("RGB")
                        _bwe, _bhe = _base_e.size
                        _swe = max(30, int(_bwe * _ew / 100))
                        _xe  = int(_bwe * _ex / 100)
                        _ye  = int(_bhe * _ey / 100)
                        st.image(overlay_sig(_base_e, _sig_e, _xe, _ye, _swe),
                                 use_container_width=True,
                                 caption="Kéo slider → xem trước cập nhật ngay")
                    except Exception as _ep:
                        st.warning(f"Không xem trước được: {_ep}")

                    if st.button("💾 Lưu vị trí này vào file",
                                 type="primary", use_container_width=True, key="b_esave"):
                        try:
                            if ext_p == ".pdf":
                                _doc_e = __import__("fitz").open(stream=_orig_b, filetype="pdf")
                                _all_pages_e = list(range(1, _doc_e.page_count + 1))
                                _new_b = sign_pdf_at_xy_pct(_orig_b, _sig_e, _all_pages_e,
                                                            _ex, _ey, _ew)
                            else:
                                _img_e = Image.open(io.BytesIO(_orig_b)).convert("RGB")
                                _bw2, _bh2 = _img_e.size
                                _sw2 = max(30, int(_bw2 * _ew / 100))
                                _res_e = overlay_sig(_img_e, _sig_e,
                                                     int(_bw2 * _ex / 100),
                                                     int(_bh2 * _ey / 100), _sw2)
                                _buf_e = io.BytesIO(); _res_e.save(_buf_e, "PNG")
                                _new_b = _buf_e.getvalue()

                            for _r_e in st.session_state.batch_results:
                                if _r_e["name"] == sel_r["name"]:
                                    _r_e["bytes"] = _new_b
                                    _r_e["detail"] += " [đã chỉnh vị trí]"
                                    break
                            st.success("Đã cập nhật! Tải file ZIP bên dưới để lấy bản mới.")
                            st.rerun()
                        except Exception as _ee:
                            st.error(f"Lỗi: {_ee}")

            # ── Tải về ───────────────────────────────────────────────────
            st.divider()
            # CSS làm nút/link tải nổi bật
            st.markdown("""<style>
                [data-testid="stDownloadButton"] button {
                    background-color: #22c55e !important;
                    color: white !important;
                    font-weight: bold !important;
                    font-size: 1.1em !important;
                    border: none !important;
                    padding: 0.6em 1em !important;
                }
                .dl-link {
                    display: block; text-align: center; padding: 14px;
                    background: #22c55e; color: white !important;
                    border-radius: 8px; font-weight: bold; font-size: 1.1em;
                    text-decoration: none; margin: 6px 0;
                }
                .dl-link:active { background: #15803d; }
                .dl-link-small {
                    display: inline-block; padding: 8px 16px;
                    background: #3b82f6; color: white !important;
                    border-radius: 6px; text-decoration: none;
                    font-size: 0.9em; margin: 4px 0;
                }
            </style>""", unsafe_allow_html=True)

            import base64
            zip_bytes = create_zip(ok)
            _b64_zip = base64.b64encode(zip_bytes).decode()

            # Link tải chính (hoạt động trên cả iOS PWA và PC)
            st.markdown(
                f'<a class="dl-link" href="data:application/zip;base64,{_b64_zip}" '
                f'download="da_ky_hang_loat.zip">⬇️ Tải tất cả {len(ok)} file đã ký (ZIP)</a>',
                unsafe_allow_html=True)

            st.info(
                "📱 **iPhone**: Bấm link trên → chọn **Tải về** hoặc **Mở bằng Files**.\n\n"
                "💻 **Máy tính**: File tự tải vào thư mục **Downloads**."
            )

            with st.expander("Tải từng file riêng"):
                for r in ok:
                    ext = _ext(r["name"])
                    _mime = ("application/pdf" if ext == ".pdf" else
                            "image/png" if ext in (".png",".jpg",".jpeg") else
                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document" if ext == ".docx" else
                            "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
                    _fname = r.get("out_name", r["name"])
                    _b64 = base64.b64encode(r["bytes"]).decode()
                    st.markdown(
                        f'<a class="dl-link-small" href="data:{_mime};base64,{_b64}" '
                        f'download="{_fname}">⬇️ {_fname}</a>',
                        unsafe_allow_html=True)
